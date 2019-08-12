# -*- coding: utf-8 -*-
"""
Created on Tue Jun 18 11:33:12 2019

@author: bhavy
"""
from docx import Document
def No_Pages_Tables_Figs(t):
    
    document = Document(t)
    section = document.sections[0]
    header = section.header
    for paragraph in header.paragraphs:
        No_of_pages=paragraph.text
    
    
    No_of_Tables=len(document.tables)
    No_of_Figure=len(document.inline_shapes)
    
    if No_of_pages=="":
        error="Cant calculate number of pages."
        return No_of_Tables,No_of_Figure,error
    else:
        return No_of_Tables,No_of_Figure,No_of_pages
        